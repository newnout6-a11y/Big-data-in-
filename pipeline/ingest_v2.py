"""Инкрементальный ингест документов в новую схему payload.

Особенности vs старый ingest.py:
- Подхватывает метаданные из harvested_meta/<hash>.json (если есть).
- Авторазметка по таксономии через классификатор.py (без LLM).
- Дедупликация чанков по text_hash.
- Расширенный payload: domain, subdomain, source, year, language,
  title, authors, journal, doi, embed_model, text_hash.
- Инкрементальный режим: по умолчанию пропускает уже обработанные документы.

Использование:
    python ingest_v2.py [--full]   # --full = переобработать всё с нуля
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import sys
import time
from datetime import datetime, timezone
from pathlib import Path

# Bootstrap: скрипт может запускаться напрямую (`python pipeline/ingest_v2.py`).
# Streamlit/python кладут в sys.path только папку самого скрипта (pipeline/),
# поэтому достраиваем sys.path до core/ и корня репо.
_РЕПО = Path(__file__).resolve().parent.parent
for _подпапка in ("", "core"):
    _путь = str(_РЕПО / _подпапка) if _подпапка else str(_РЕПО)
    if _путь not in sys.path:
        sys.path.insert(0, _путь)

import docx

try:
    import fitz  # PyMuPDF — быстрее и лучше с шрифтами/формулами
    _ИСПОЛЬЗОВАТЬ_PYMUPDF = True
except ImportError:
    fitz = None
    _ИСПОЛЬЗОВАТЬ_PYMUPDF = False

import pypdf  # fallback и для не-PyMuPDF окружений

from cases import определить_кейс
from классификатор import (
    детерминировать_язык,
    классифицировать_батч,
    подготовить_прототипы,
)
import фильтр_качества
import извлечение_картинок as извк


# Скрипт лежит в pipeline/, входные/выходные файлы — в корне репо.
_БАЗА = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
ПАПКА_PDF = os.path.join(_БАЗА, "all_pdfs")
ПАПКА_МЕТА = os.path.join(_БАЗА, "harvested_meta")
ПАПКА_ОТБРАКОВКА = os.path.join(_БАЗА, "rejected_pdfs")
ПАПКА_КАРТИНОК = os.path.join(_БАЗА, "extracted_images")
ФАЙЛ_ЧАНКОВ = os.path.join(_БАЗА, "chunks_v2.jsonl")

РАЗМЕР_ЧАНКА = 800
ПЕРЕКРЫТИЕ = 100
СИМВ_НА_СТРАНИЦУ_DOCX = 2500
EMBED_MODEL_TAG = "e5-base-v1"


def _хэш_файла(путь):
    h = hashlib.sha256()
    with open(путь, "rb") as f:
        for блок in iter(lambda: f.read(1024 * 1024), b""):
            h.update(блок)
    return h.hexdigest()


def _относительный_путь(путь):
    return os.path.relpath(путь, _БАЗА).replace(os.sep, "/")


def _извлечь_картинки_страницы(документ, страница, номер_страницы, папка_документа, фоновые_xref=None):
    """Извлекает картинки страницы с полным набором обработки:
    фильтр декоративных, bbox-привязка подписей Fig./Рис./Scheme/Table,
    fallback-caption из заголовка слайда.

    В harvester pipeline (HARVEST_SKIP_IMAGES=1) картинки пропускаются,
    чтобы не раздувать Drive.
    """
    if os.getenv("HARVEST_SKIP_IMAGES") == "1":
        return []
    текст_страницы = ""
    try:
        текст_страницы = страница.get_text("text") or ""
    except Exception:
        pass
    return извк.извлечь_картинки_страницы(
        документ,
        страница,
        номер_страницы,
        Path(папка_документа),
        текст_страницы=текст_страницы,
        фоновые_xref=фоновые_xref,
        use_ocr=False,
        base_dir=Path(_БАЗА),
    )


def _извлечь_pdf_pymupdf(путь):
    страницы = []
    картинки_по_страницам = {}
    try:
        хэш_pdf = _хэш_файла(путь)
        папка_документа = os.path.join(ПАПКА_КАРТИНОК, хэш_pdf)
        документ = fitz.open(путь)
        try:
            фоновые = извк.найти_фоновые_xref(документ)
            for номер, страница in enumerate(документ, start=1):
                try:
                    текст = страница.get_text("text") or ""
                except Exception:
                    continue
                картинки = _извлечь_картинки_страницы(
                    документ, страница, номер, папка_документа, фоновые_xref=фоновые
                )
                if картинки:
                    картинки_по_страницам[номер] = картинки
                if текст and len(текст.strip()) > 50:
                    страницы.append((номер, текст.strip()))
        finally:
            документ.close()
    except Exception as ошибка:
        print(f"  ОШИБКА PyMuPDF: {ошибка}")
    return страницы, картинки_по_страницам


def _извлечь_pdf_pypdf(путь):
    страницы = []
    try:
        читалка = pypdf.PdfReader(путь)
        for номер, страница in enumerate(читалка.pages, start=1):
            try:
                текст = страница.extract_text()
            except Exception:
                continue
            if текст and len(текст.strip()) > 50:
                страницы.append((номер, текст.strip()))
    except Exception as ошибка:
        print(f"  ОШИБКА pypdf: {ошибка}")
    return страницы, {}


def извлечь_pdf(путь):
    """PyMuPDF когда доступен (5–10× быстрее, лучше с шрифтами), иначе pypdf."""
    if _ИСПОЛЬЗОВАТЬ_PYMUPDF:
        страницы, картинки_по_страницам = _извлечь_pdf_pymupdf(путь)
        if страницы:
            return страницы, картинки_по_страницам
        # Если PyMuPDF ничего не извлёк (защищённый/повреждённый PDF), пробуем pypdf
    return _извлечь_pdf_pypdf(путь)


def извлечь_docx(путь):
    try:
        документ = docx.Document(путь)
        куски = [п.text.strip() for п in документ.paragraphs if п.text.strip()]
        for таблица in документ.tables:
            for строка in таблица.rows:
                ряд = " | ".join(я.text.strip() for я in строка.cells if я.text.strip())
                if ряд:
                    куски.append(ряд)
        полный = "\n".join(куски)
    except Exception as ошибка:
        print(f"  ОШИБКА docx: {ошибка}")
        return []

    if len(полный.strip()) < 50:
        return []

    страницы = []
    номер = 1
    начало = 0
    while начало < len(полный):
        конец = начало + СИМВ_НА_СТРАНИЦУ_DOCX
        кусок = полный[начало:конец].strip()
        if len(кусок) > 50:
            страницы.append((номер, кусок))
        начало = конец
        номер += 1
    return страницы


def извлечь_txt(путь):
    """Простой текстовый файл (например, Stack Exchange Q+A)."""
    try:
        with open(путь, "r", encoding="utf-8", errors="ignore") as f:
            полный = f.read().strip()
    except Exception as ошибка:
        print(f"  ОШИБКА txt: {ошибка}")
        return []
    if len(полный) < 50:
        return []
    страницы = []
    номер = 1
    начало = 0
    while начало < len(полный):
        конец = начало + СИМВ_НА_СТРАНИЦУ_DOCX
        кусок = полный[начало:конец].strip()
        if len(кусок) > 50:
            страницы.append((номер, кусок))
        начало = конец
        номер += 1
    return страницы


def разбить_на_чанки(текст, размер=РАЗМЕР_ЧАНКА, перекр=ПЕРЕКРЫТИЕ):
    чанки = []
    начало = 0
    while начало < len(текст):
        конец = начало + размер
        кусок = текст[начало:конец].strip()
        if len(кусок) > 80:
            чанки.append(кусок)
        начало = конец - перекр
    return чанки


def прочитать_метадату(имя_файла):
    """harvested_meta/<имя_без_расширения>.json — если положил harvester."""
    основание, _ = os.path.splitext(имя_файла)
    путь = os.path.join(ПАПКА_МЕТА, основание + ".json")
    if not os.path.exists(путь):
        return {}
    try:
        with open(путь, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def собрать_обработанные():
    обработанные = set()
    if not os.path.exists(ФАЙЛ_ЧАНКОВ):
        return обработанные
    with open(ФАЙЛ_ЧАНКОВ, "r", encoding="utf-8") as f:
        for строка in f:
            строка = строка.strip()
            if not строка:
                continue
            try:
                запись = json.loads(строка)
                обработанные.add(запись.get("document", ""))
            except Exception:
                continue
    return обработанные


def main(argv=None):
    парсер = argparse.ArgumentParser()
    парсер.add_argument("--full", action="store_true", help="Перезаписать chunks_v2.jsonl с нуля")
    парсер.add_argument("--no-quality-filter", action="store_true",
                        help="Не выкидывать низкокачественные PDF в rejected_pdfs/")
    args = парсер.parse_args(argv)

    if not os.path.isdir(ПАПКА_PDF):
        print(f"Нет папки {ПАПКА_PDF}")
        return 1

    обработанные = set() if args.full else собрать_обработанные()
    режим = "w" if args.full else "a"

    все_файлы = sorted([
        ф for ф in os.listdir(ПАПКА_PDF)
        if ф.lower().endswith((".pdf", ".docx", ".txt"))
    ])
    новые = [ф for ф in все_файлы if ф not in обработанные]
    print(f"Всего файлов: {len(все_файлы)}, новых: {len(новые)}")
    if not новые:
        print("Нечего обрабатывать.")
        return 0

    # Загрузка модели и прототипов один раз
    print("Загружаю модель эмбеддингов и прототипы…")
    from sentence_transformers import SentenceTransformer
    модель = SentenceTransformer("intfloat/multilingual-e5-base")
    метки, прототипы, _ = подготовить_прототипы(модель)

    счётчик_чанков = 0
    отбракованных = 0
    хэши_в_сессии: set[str] = set()
    с_дедупом = 0
    время_старта = time.time()

    with open(ФАЙЛ_ЧАНКОВ, режим, encoding="utf-8") as выход:
        for индекс, имя in enumerate(новые, start=1):
            t_файл = time.time()
            размер_МБ = os.path.getsize(os.path.join(ПАПКА_PDF, имя)) / (1024 * 1024)
            print(f"[{индекс}/{len(новые)}] {имя} ({размер_МБ:.1f} MB) ...", flush=True)
            путь = os.path.join(ПАПКА_PDF, имя)
            нижний = имя.lower()
            картинки_по_страницам = {}
            if нижний.endswith(".pdf"):
                страницы, картинки_по_страницам = извлечь_pdf(путь)
            elif нижний.endswith(".docx"):
                страницы = извлечь_docx(путь)
            else:
                страницы = извлечь_txt(путь)
            if not страницы:
                print(f"  [{индекс}/{len(новые)}] ПРОПУЩЕН (нет текста): {имя}")
                if not args.no_quality_filter and нижний.endswith(".pdf"):
                    отбракованных += 1
                    куда = фильтр_качества.отбраковать(путь, ПАПКА_ОТБРАКОВКА, "нет извлечённого текста")
                    if куда:
                        print(f"      → перемещён в {куда}")
                continue

            # Фильтр качества: отбрасываем низкокачественные PDF до ingest
            if not args.no_quality_filter and нижний.endswith(".pdf"):
                оценка = фильтр_качества.оценить(страницы)
                if not оценка.принят:
                    отбракованных += 1
                    print(f"  [{индекс}/{len(новые)}] ОТБРАКОВАН ({оценка.причина}): {имя}")
                    куда = фильтр_качества.отбраковать(путь, ПАПКА_ОТБРАКОВКА, оценка.причина)
                    if куда:
                        print(f"      → перемещён в {куда}")
                    continue

            мета = прочитать_метадату(имя)
            doc_id = мета.get("doc_id") or f"local:{имя}"
            источник = мета.get("источник", "local")
            название = мета.get("название", "")
            авторы = мета.get("авторы", [])
            дата_публ = мета.get("дата", "")
            год = None
            if дата_публ[:4].isdigit():
                год = int(дата_публ[:4])

            # Собираем все чанки этого файла
            пары: list[tuple[int, str, list[dict]]] = []
            for номер_страницы, текст_страницы in страницы:
                for чанк in разбить_на_чанки(текст_страницы):
                    хэш = hashlib.sha1(чанк.encode("utf-8")).hexdigest()
                    if хэш in хэши_в_сессии:
                        с_дедупом += 1
                        continue
                    хэши_в_сессии.add(хэш)
                    пары.append((
                        номер_страницы,
                        чанк,
                        list(картинки_по_страницам.get(номер_страницы, [])),
                    ))

            if not пары:
                print(f"  [{индекс}/{len(новые)}] ПУСТО после дедупа: {имя}")
                continue

            # Авторазметка батчем
            тексты = [ч for _, ч, _ in пары]
            метки_чанков = классифицировать_батч(тексты, модель, метки, прототипы)

            for (номер_стр, чанк, картинки), (домен, суб, скор) in zip(пары, метки_чанков):
                язык = детерминировать_язык(чанк)
                кейс = определить_кейс(чанк)  # для backward-compat
                хэш_текста = hashlib.sha1(чанк.encode("utf-8")).hexdigest()
                запись = {
                    "text": чанк,
                    "document": имя,
                    "page": номер_стр,
                    "images": картинки,
                    "case": кейс,                       # backward-compat
                    "doc_id": doc_id,
                    "source": источник,
                    "title": название,
                    "authors": авторы,
                    "year": год,
                    "domain": домен,
                    "subdomain": суб,
                    "topic_score": round(скор, 3),
                    "language": язык,
                    "embed_model": EMBED_MODEL_TAG,
                    "text_hash": хэш_текста,
                    "ingested_at": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
                }
                выход.write(json.dumps(запись, ensure_ascii=False) + "\n")
                счётчик_чанков += 1
            выход.flush()

            длит = time.time() - t_файл
            прошло = time.time() - время_старта
            сред = прошло / индекс
            осталось = (len(новые) - индекс) * сред
            print(
                f"  → {длит:.1f}s, chunks={счётчик_чанков}, прошло {прошло/60:.1f}m, ETA ≈ {осталось/60:.1f}m",
                flush=True,
            )

    print(f"\nГотово. Чанков добавлено: {счётчик_чанков}, дедупликаций: {с_дедупом}")
    if отбракованных:
        print(f"Отбраковано (low quality): {отбракованных} → {ПАПКА_ОТБРАКОВКА}/")
    print(f"Файл: {ФАЙЛ_ЧАНКОВ}")
    print("Дальше: python embed_resume_v2.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
