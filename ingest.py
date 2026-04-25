import os
import json
import pypdf
import docx
from cases import определить_кейс

папка_с_пдф = os.path.join(os.path.dirname(os.path.abspath(__file__)), "all_pdfs")
файл_чанков = os.path.join(os.path.dirname(os.path.abspath(__file__)), "chunks.jsonl")

размер_чанка = 800
перекрытие = 100
симв_на_страницу_docx = 2500


def извлечь_текст_из_пдф(путь):
    страницы = []
    try:
        читалка = pypdf.PdfReader(путь)
        for номер, страница in enumerate(читалка.pages, start=1):
            текст = страница.extract_text()
            if текст and len(текст.strip()) > 50:
                страницы.append((номер, текст.strip()))
    except Exception:
        pass
    return страницы


def извлечь_текст_из_docx(путь):
    try:
        документ = docx.Document(путь)
        куски = []
        for пар in документ.paragraphs:
            т = пар.text.strip()
            if т:
                куски.append(т)
        for таблица in документ.tables:
            for строка in таблица.rows:
                ряд = " | ".join(ячейка.text.strip() for ячейка in строка.cells if ячейка.text.strip())
                if ряд:
                    куски.append(ряд)
        полный = "\n".join(куски)
    except Exception:
        return []

    if len(полный.strip()) < 50:
        return []

    страницы = []
    номер = 1
    начало = 0
    while начало < len(полный):
        конец = начало + симв_на_страницу_docx
        кусок = полный[начало:конец].strip()
        if len(кусок) > 50:
            страницы.append((номер, кусок))
        начало = конец
        номер += 1
    return страницы


def разбить_на_чанки(текст, размер, перекр):
    чанки = []
    начало = 0
    while начало < len(текст):
        конец = начало + размер
        кусок = текст[начало:конец].strip()
        if len(кусок) > 80:
            чанки.append(кусок)
        начало = конец - перекр
    return чанки


все_файлы = sorted([
    ф for ф in os.listdir(папка_с_пдф)
    if ф.lower().endswith((".pdf", ".docx"))
])

print(f"Найдено файлов: {len(все_файлы)}")

счётчик_чанков = 0
счётчик_файлов = 0

with open(файл_чанков, "w", encoding="utf-8") as выходной:
    for индекс, имя_файла in enumerate(все_файлы, start=1):
        полный_путь = os.path.join(папка_с_пдф, имя_файла)
        if имя_файла.lower().endswith(".pdf"):
            страницы = извлечь_текст_из_пдф(полный_путь)
        else:
            страницы = извлечь_текст_из_docx(полный_путь)

        if not страницы:
            print(f"  [{индекс}/{len(все_файлы)}] ПРОПУЩЕН (нет текста): {имя_файла}")
            continue

        for номер_страницы, текст_страницы in страницы:
            чанки = разбить_на_чанки(текст_страницы, размер_чанка, перекрытие)
            for чанк in чанки:
                кейс = определить_кейс(чанк)
                запись = {
                    "text": чанк,
                    "document": имя_файла,
                    "page": номер_страницы,
                    "case": кейс
                }
                выходной.write(json.dumps(запись, ensure_ascii=False) + "\n")
                счётчик_чанков += 1

        счётчик_файлов += 1
        if индекс % 50 == 0 or индекс == len(все_файлы):
            print(f"  [{индекс}/{len(все_файлы)}] обработано файлов: {счётчик_файлов}, чанков: {счётчик_чанков}")

print(f"\nГотово! Всего чанков: {счётчик_чанков}")
print(f"Сохранено в: {файл_чанков}")
