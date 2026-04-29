"""Инкрементальный ингест документов в новую схему payload.

Особенности vs старый ingest.py:
- Подхватывает метаданные из harvested_meta/<hash>.json (если есть).
- Авторазметка по таксономии через классификатор.py (без LLM).
- Дедупликация чанков по text_hash.
- Расширенный payload: domain, subdomain, source, year, language,
  title, authors, journal, doi, embed_model, text_hash.
- Инкрементальный режим: по умолчанию пропускает уже обработанные документы.

Использование:
    python ingest_v2.py [--full]   # --full = переобработать всё с нуля
"""
from __future__ import annotations

import argparse
import hashlib
import json
import os
import sys
from datetime import datetime

import docx

try:
    import fitz  # PyMuPDF — быстрее и лучше с шрифтами/формулами
    _ИСПОЛЬЗОВАТЬ_PYMUPDF = True
except ImportError:
    fitz = None
    _ИСПОЛЬЗОВАТЬ_PYMUPDF = False

import pypdf  # fallback и для не-PyMuPDF окружений

from cases import определить_кейс
from классификатор import (
    детерминировать_язык,
    классифицировать_батч,
    подготовить_прототипы,
)


_БАЗА = os.path.dirname(os.path.abspath(__file__))
ПАПКА_PDF = os.path.join(_БАЗА, "all_pdfs")
ПАПКА_МЕТА = os.path.join(_БАЗА, "harvested_meta")
ФАЙЛ_ЧАНКОВ = os.path.join(_БАЗА, "chunks_v2.jsonl")

РАЗМЕР_ЧАНКА = 800
ПЕРЕКРЫТИЕ = 100
СИМВ_НА_СТРАНИЦУ_DOCX = 2500
EMBED_MODEL_TAG = "e5-base-v1"


def _извлечь_pdf_pymupdf(путь):
    страницы = []
    try:
        документ = fitz.open(путь)
        try:
            for номер, страница in enumerate(документ, start=1):
                try:
                    текст = страница.get_text("text") or ""
                except Exception:
                    continue
                if текст and len(текст.strip()) > 50:
                    страницы.append((номер, текст.strip()))
        finally:
            документ.close()
    except Exception as ошибка:
        print(f"  ОШИБКА PyMuPDF: {ошибка}")
    return страницы


def _извлечь_pdf_pypdf(путь):
    страницы = []
    try:
        читалка = pypdf.PdfReader(путь)
        for номер, страница in enumerate(читалка.pages, start=1):
            try:
                текст = страница.extract_text()
            except Exception:
                continue
            if текст and len(текст.strip()) > 50:
                страницы.append((номер, текст.strip()))
    except Exception as ошибка:
        print(f"  ОШИБКА pypdf: {ошибка}")
    return страницы


def извлечь_pdf(путь):
    """PyMuPDF когда доступен (5–10× быстрее, лучше с шрифтами), иначе pypdf."""
    if _ИСПОЛЬЗОВАТЬ_PYMUPDF:
        страницы = _извлечь_pdf_pymupdf(путь)
        if страницы:
            return страницы
        # Если PyMuPDF ничего не извлёк (защищённый/повреждённый PDF), пробуем pypdf
    return _извлечь_pdf_pypdf(путь)


def извлечь_docx(путь):
    try:
        документ = docx.Document(путь)
        куски = [п.text.strip() for п in документ.paragraphs if п.text.strip()]
        for таблица in документ.tables:
            for строка in таблица.rows:
                ряд = " | ".join(я.text.strip() for я in строка.cells if я.text.strip())
                if ряд:
                    куски.append(ряд)
        полный = "\n".join(куски)
    except Exception as ошибка:
        print(f"  ОШИБКА docx: {ошибка}")
        return []

    if len(полный.strip()) < 50:
        return []

    страницы = []
    номер = 1
    начало = 0
    while начало < len(полный):
        конец = начало + СИМВ_НА_СТРАНИЦУ_DOCX
        кусок = полный[начало:конец].strip()
        if len(кусок) > 50:
            страницы.append((номер, кусок))
        начало = конец
        номер += 1
    return страницы


def извлечь_txt(путь):
    """Простой текстовый файл (например, Stack Exchange Q+A)."""
    try:
        with open(путь, "r", encoding="utf-8", errors="ignore") as f:
            полный = f.read().strip()
    except Exception as ошибка:
        print(f"  ОШИБКА txt: {ошибка}")
        return []
    if len(полный) < 50:
        return []
    страницы = []
    номер = 1
    начало = 0
    while начало < len(полный):
        конец = начало + СИМВ_НА_СТРАНИЦУ_DOCX
        кусок = полный[начало:конец].strip()
        if len(кусок) > 50:
            страницы.append((номер, кусок))
        начало = конец
        номер += 1
    return страницы


def разбить_на_чанки(текст, размер=РАЗМЕР_ЧАНКА, перекр=ПЕРЕКРЫТИЕ):
    чанки = []
    начало = 0
    while начало < len(текст):
        конец = начало + размер
        кусок = текст[начало:конец].strip()
        if len(кусок) > 80:
            чанки.append(кусок)
        начало = конец - перекр
    return чанки


def прочитать_метадату(имя_файла):
    """harvested_meta/<имя_без_расширения>.json — если положил harvester."""
    основание, _ = os.path.splitext(имя_файла)
    путь = os.path.join(ПАПКА_МЕТА, основание + ".json")
    if not os.path.exists(путь):
        return {}
    try:
        with open(путь, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def собрать_обработанные():
    обработанные = set()
    if not os.path.exists(ФАЙЛ_ЧАНКОВ):
        return обработанные
    with open(ФАЙЛ_ЧАНКОВ, "r", encoding="utf-8") as f:
        for строка in f:
            строка = строка.strip()
            if not строка:
                continue
            try:
                запись = json.loads(строка)
                обработанные.add(запись.get("document", ""))
            except Exception:
                continue
    return обработанные


def main(argv=None):
    парсер = argparse.ArgumentParser()
    парсер.add_argument("--full", action="store_true", help="Перезаписать chunks_v2.jsonl с нуля")
    args = парсер.parse_args(argv)

    if not os.path.isdir(ПАПКА_PDF):
        print(f"Нет папки {ПАПКА_PDF}")
        return 1

    обработанные = set() if args.full else собрать_обработанные()
    режим = "w" if args.full else "a"

    все_файлы = sorted([
        ф for ф in os.listdir(ПАПКА_PDF)
        if ф.lower().endswith((".pdf", ".docx", ".txt"))
    ])
    новые = [ф for ф in все_файлы if ф not in обработанные]
    print(f"Всего файлов: {len(все_файлы)}, новых: {len(новые)}")
    if not новые:
        print("Нечего обрабатывать.")
        return 0

    # Загрузка модели и прототипов один раз
    print("Загружаю модель эмбеддингов и прототипы…")
    from sentence_transformers import SentenceTransformer
    модель = SentenceTransformer("intfloat/multilingual-e5-base")
    метки, прототипы, _ = подготовить_прототипы(модель)

    счётчик_чанков = 0
    хэши_в_сессии: set[str] = set()
    с_дедупом = 0

    with open(ФАЙЛ_ЧАНКОВ, режим, encoding="utf-8") as выход:
        for индекс, имя in enumerate(новые, start=1):
            путь = os.path.join(ПАПКА_PDF, имя)
            нижний = имя.lower()
            if нижний.endswith(".pdf"):
                страницы = извлечь_pdf(путь)
            elif нижний.endswith(".docx"):
                страницы = извлечь_docx(путь)
            else:
                страницы = извлечь_txt(путь)
            if not страницы:
                print(f"  [{индекс}/{len(новые)}] ПРОПУЩЕН (нет текста): {имя}")
                continue

            мета = прочитать_метадату(имя)
            doc_id = мета.get("doc_id") or f"local:{имя}"
            источник = мета.get("источник", "local")
            название = мета.get("название", "")
            авторы = мета.get("авторы", [])
            дата_публ = мета.get("дата", "")
            год = None
            if дата_публ[:4].isdigit():
                год = int(дата_публ[:4])

            # Собираем все чанки этого файла
            пары: list[tuple[int, str]] = []
            for номер_страницы, текст_страницы in страницы:
                for чанк in разбить_на_чанки(текст_страницы):
                    хэш = hashlib.sha1(чанк.encode("utf-8")).hexdigest()
                    if хэш in хэши_в_сессии:
                        с_дедупом += 1
                        continue
                    хэши_в_сессии.add(хэш)
                    пары.append((номер_страницы, чанк))

            if not пары:
                print(f"  [{индекс}/{len(новые)}] ПУСТО после дедупа: {имя}")
                continue

            # Авторазметка батчем
            тексты = [ч for _, ч in пары]
            метки_чанков = классифицировать_батч(тексты, модель, метки, прототипы)

            for (номер_стр, чанк), (домен, суб, скор) in zip(пары, метки_чанков):
                язык = детерминировать_язык(чанк)
                кейс = определить_кейс(чанк)  # для backward-compat
                запись = {
                    "text": чанк,
                    "document": имя,
                    "page": номер_стр,
                    "case": кейс,                       # backward-compat
                    "doc_id": doc_id,
                    "source": источник,
                    "title": название,
                    "authors": авторы,
                    "year": год,
                    "domain": домен,
                    "subdomain": суб,
                    "topic_score": round(скор, 3),
                    "language": язык,
                    "embed_model": EMBED_MODEL_TAG,
                    "text_hash": hashlib.sha1(чанк.encode("utf-8")).hexdigest(),
                    "ingested_at": datetime.utcnow().strftime("%Y-%m-%d"),
                }
                выход.write(json.dumps(запись, ensure_ascii=False) + "\n")
                счётчик_чанков += 1
            выход.flush()

            if индекс % 25 == 0 or индекс == len(новые):
                print(f"  [{индекс}/{len(новые)}] чанков добавлено всего: {счётчик_чанков}")

    print(f"\nГотово. Чанков добавлено: {счётчик_чанков}, дедупликаций: {с_дедупом}")
    print(f"Файл: {ФАЙЛ_ЧАНКОВ}")
    print("Дальше: python embed_resume_v2.py")
    return 0


if __name__ == "__main__":
    sys.exit(main())
