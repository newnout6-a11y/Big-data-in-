from __future__ import annotations

import base64
import hashlib
import io
import json
import mimetypes
import re
import tempfile
import csv
from datetime import datetime
from pathlib import Path
from typing import Any, Callable

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.shared import Inches, Pt


# Маркер картинки в ответе LLM: [img:N.M] — номер фрагмента и номер картинки
# внутри него. Тот же формат используется в app.py:_IMG_МАРКЕР.
_IMG_МАРКЕР = re.compile(r"\[img:(\d+)\.(\d+)\]")


# Тип функции-резолвера: по (N, M) вернуть (путь, подпись, caption) или None.
# Сигнатура совпадает с app._разрешить_img_маркер, но здесь она опциональна,
# чтобы study_tools оставался не-привязанным к app.
ImageResolver = Callable[[int, int], "tuple[Path, str, str] | None"]


def fragments_context(fragments: list[dict[str, Any]], *, max_chars: int = 12000) -> str:
    parts = []
    used = 0
    for index, fragment in enumerate(fragments, 1):
        text = str(fragment.get("text", "")).strip()
        if not text:
            continue
        header = f"[{index}] Документ: {fragment.get('document', '')}, стр. {fragment.get('page', '')}\n"
        block = header + text + "\n\n"
        if used + len(block) > max_chars and parts:
            break
        parts.append(block)
        used += len(block)
    return "".join(parts)


def parse_json_loose(raw: str) -> dict[str, Any]:
    text = (raw or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    try:
        data = json.loads(text)
        return data if isinstance(data, dict) else {"items": data}
    except json.JSONDecodeError:
        pass

    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        data = json.loads(text[start:end + 1])
        return data if isinstance(data, dict) else {"items": data}
    raise ValueError("LLM вернула невалидный JSON.")


def _картинка_data_uri(путь: Path) -> str | None:
    """Читает картинку и возвращает data: URI, пригодный для встраивания
    в Markdown. Возвращает None если файл не читается."""
    try:
        data = путь.read_bytes()
    except (OSError, ValueError):
        return None
    mime, _ = mimetypes.guess_type(str(путь))
    if not mime:
        mime = "image/png"
    b64 = base64.b64encode(data).decode("ascii")
    return f"data:{mime};base64,{b64}"


def markdown_export(
    title: str,
    body: str,
    fragments: list[dict[str, Any]] | None = None,
    *,
    image_resolver: ImageResolver | None = None,
    embed_base64: bool = True,
) -> bytes:
    """Собирает .md из тела ответа. Маркеры [img:N.M] заменяет на
    Markdown-картинки: если `embed_base64=True` — base64-data-URI (файл
    самодостаточен и не требует папки с картинками рядом), иначе —
    относительный путь относительно текущей папки.

    LaTeX-блоки $$...$$ и $...$ оставляем как есть — современные
    Markdown-рендеры (Obsidian, Typora, GitHub с MathJax) их понимают.
    Артефакты вроде голых `\\text{...}` вне $-блоков чистим: не каждый
    рендер умеет, и без $ они смотрятся как мусор.
    """
    body = _очистить_latex_вне_формул(body.strip())
    body = _заменить_img_маркеры_на_markdown(
        body,
        image_resolver=image_resolver,
        embed_base64=embed_base64,
    )

    lines = [f"# {title.strip() or 'Экспорт'}", "", body, ""]
    if fragments:
        lines.extend(["## Источники", ""])
        for index, fragment in enumerate(fragments, 1):
            lines.append(
                f"- [{index}] {fragment.get('document', '')}, стр. {fragment.get('page', '')}"
            )
    return ("\n".join(lines).strip() + "\n").encode("utf-8")


def docx_export(
    title: str,
    body: str,
    fragments: list[dict[str, Any]] | None = None,
    *,
    image_resolver: ImageResolver | None = None,
) -> bytes:
    """Собирает .docx из тела ответа. Маркеры [img:N.M] резолвятся через
    `image_resolver` и встраиваются как реальные картинки с подписью.
    Без резолвера — маркер просто удаляется (в документе не остаётся
    сырого `[img:1.2]`).

    LaTeX: блоки $$...$$ и inline $...$ рендерятся в OMML, включая
    `\\text{}`/`\\mathrm{}` (Word показывает обычным шрифтом внутри
    формулы). Голый `\\text{...}` вне $-блоков нормализуется до просто
    его содержимого — иначе в документе остаются артефакты.
    """
    document = docx.Document()
    document.add_heading(title.strip() or "Экспорт", level=1)

    body = body.strip()

    # Режем тело по маркерам картинок. Между маркерами — обычный текст,
    # который передаём в старый парсер ($$ / $ / заголовки / параграфы).
    последняя_позиция = 0
    for m in _IMG_МАРКЕР.finditer(body):
        кусок = body[последняя_позиция:m.start()]
        if кусок.strip():
            _docx_добавить_текстовый_блок(document, кусок)
        последняя_позиция = m.end()
        if image_resolver is not None:
            _docx_добавить_картинку(document, int(m.group(1)), int(m.group(2)), image_resolver)
        # Если резолвера нет или картинка не нашлась — просто пропускаем маркер.

    хвост = body[последняя_позиция:]
    if хвост.strip():
        _docx_добавить_текстовый_блок(document, хвост)

    if fragments:
        document.add_heading("Источники", level=2)
        for index, fragment in enumerate(fragments, 1):
            document.add_paragraph(
                f"[{index}] {fragment.get('document', '')}, стр. {fragment.get('page', '')}"
            )

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _docx_добавить_текстовый_блок(document: Any, часть: str) -> None:
    """Парсит произвольный кусок текста (без маркеров картинок) и добавляет
    в document соответствующие paragraph'ы/формулы."""
    часть = _очистить_latex_вне_формул(часть.strip())
    if not часть:
        return
    for part in re.split(r"(\$\$.*?\$\$)", часть, flags=re.DOTALL):
        part = part.strip()
        if not part:
            continue
        if part.startswith("$$") and part.endswith("$$"):
            latex = part[2:-2].strip()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph._p.append(_omath(latex))
            continue
        for block in re.split(r"\n{2,}", part):
            block = block.strip()
            if not block:
                continue
            if block.startswith(("# ", "## ", "### ")):
                level = min(block.count("#", 0, block.find(" ")), 3)
                document.add_heading(block.lstrip("# ").strip(), level=level)
            else:
                paragraph = document.add_paragraph()
                _append_inline_math(paragraph, block)


def _docx_добавить_картинку(
    document: Any,
    n: int,
    m: int,
    image_resolver: ImageResolver,
) -> None:
    try:
        разрешено = image_resolver(n, m)
    except Exception:
        разрешено = None
    if разрешено is None:
        return
    путь, подпись, _caption = разрешено
    try:
        document.add_picture(str(путь), width=Inches(5.5))
    except Exception:
        # Битый/неподдерживаемый файл — не валим весь экспорт, просто
        # добавим подпись-плейсхолдер.
        p = document.add_paragraph()
        run = p.add_run(f"[картинка недоступна] {подпись}")
        run.italic = True
        return
    # Картинка добавлена — последним параграф выравниваем по центру,
    # следом — подпись.
    try:
        document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass
    подпись_p = document.add_paragraph()
    подпись_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = подпись_p.add_run(подпись)
    run.italic = True
    run.font.size = Pt(9)


_TEXT_КОМАНДА = re.compile(r"\\(?:text|mathrm|textbf|textit)\s*\{([^{}]*)\}")


def _очистить_latex_вне_формул(текст: str) -> str:
    """Убирает голые LaTeX-команды `\\text{...}` / `\\mathrm{...}` вне
    $-блоков: заменяет на их содержимое. Внутри $...$ и $$...$$ не трогаем —
    там их корректно рендерит OMML-парсер."""
    if not текст:
        return текст

    куски: list[str] = []
    i = 0
    while i < len(текст):
        # Ищем ближайший $-блок
        dollar = текст.find("$", i)
        if dollar == -1:
            куски.append(_TEXT_КОМАНДА.sub(lambda m: m.group(1), текст[i:]))
            break
        # Текст до блока — чистим
        куски.append(_TEXT_КОМАНДА.sub(lambda m: m.group(1), текст[i:dollar]))
        # Определяем, $$ это или $
        if текст.startswith("$$", dollar):
            end = текст.find("$$", dollar + 2)
            if end == -1:
                куски.append(текст[dollar:])
                break
            куски.append(текст[dollar:end + 2])
            i = end + 2
        else:
            end = текст.find("$", dollar + 1)
            if end == -1:
                куски.append(текст[dollar:])
                break
            куски.append(текст[dollar:end + 1])
            i = end + 1
    return "".join(куски)


def _заменить_img_маркеры_на_markdown(
    текст: str,
    *,
    image_resolver: ImageResolver | None,
    embed_base64: bool,
) -> str:
    def _замена(m: re.Match) -> str:
        if image_resolver is None:
            return ""
        try:
            разрешено = image_resolver(int(m.group(1)), int(m.group(2)))
        except Exception:
            разрешено = None
        if разрешено is None:
            return ""
        путь, подпись, _caption = разрешено
        alt = подпись.replace("[", "(").replace("]", ")")
        if embed_base64:
            data_uri = _картинка_data_uri(путь)
            if data_uri is None:
                return ""
            return f"\n\n![{alt}]({data_uri})\n\n*{alt}*\n\n"
        # Относительный путь относительно CWD — для случая, когда .md
        # распаковывается рядом с папкой extracted_images/.
        try:
            rel = путь.resolve().as_posix()
        except Exception:
            rel = str(путь)
        return f"\n\n![{alt}]({rel})\n\n*{alt}*\n\n"

    return _IMG_МАРКЕР.sub(_замена, текст)


def _append_inline_math(paragraph: Any, text: str) -> None:
    for part in re.split(r"(\$[^$\n]+\$)", text):
        if not part:
            continue
        if part.startswith("$") and part.endswith("$"):
            paragraph._p.append(_omath(part[1:-1].strip()))
        else:
            paragraph.add_run(part)


def _omath(latex: str) -> Any:
    xml = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + _math_fragments(_normalize_latex(latex)) +
        "</m:oMath>"
    )
    return parse_xml(xml)


def _math_fragments(expr: str) -> str:
    out = []
    i = 0
    while i < len(expr):
        if expr.startswith(r"\frac", i):
            i += len(r"\frac")
            num, i = _read_group(expr, i)
            den, i = _read_group(expr, i)
            out.append(
                "<m:f><m:fPr><m:type m:val=\"bar\"/></m:fPr>"
                f"<m:num>{_math_fragments(num)}</m:num>"
                f"<m:den>{_math_fragments(den)}</m:den>"
                "</m:f>"
            )
            continue

        # \text{...} / \mathrm{...} / \textbf{...} / \textit{...}: содержимое
        # вставляется как обычный run (в формуле Word покажет тем же шрифтом,
        # но без курсива). Без этого LLM-ответы с \text{Схема отстойника:}
        # остаются сырыми в docx — именно эту багу ловим.
        for команда in (r"\text", r"\mathrm", r"\textbf", r"\textit"):
            if expr.startswith(команда, i) and i + len(команда) < len(expr) \
                    and expr[i + len(команда)] == "{":
                j = i + len(команда)
                содержимое, i = _read_group(expr, j)
                out.append(_math_text(содержимое))
                break
        else:
            # Не сработало ни одно `\text*` — идём в обычную ветку
            token, i_new = _read_token(expr, i)
            if not token:
                break
            sub = sup = None
            i = i_new
            while i < len(expr) and expr[i] in "_^":
                marker = expr[i]
                value, i = _read_script(expr, i + 1)
                if marker == "_":
                    sub = value
                else:
                    sup = value
            if sub is not None and sup is not None:
                out.append(
                    "<m:sSubSup>"
                    f"<m:e>{_math_text(token)}</m:e>"
                    f"<m:sub>{_math_fragments(sub)}</m:sub>"
                    f"<m:sup>{_math_fragments(sup)}</m:sup>"
                    "</m:sSubSup>"
                )
            elif sub is not None:
                out.append(
                    "<m:sSub>"
                    f"<m:e>{_math_text(token)}</m:e>"
                    f"<m:sub>{_math_fragments(sub)}</m:sub>"
                    "</m:sSub>"
                )
            elif sup is not None:
                out.append(
                    "<m:sSup>"
                    f"<m:e>{_math_text(token)}</m:e>"
                    f"<m:sup>{_math_fragments(sup)}</m:sup>"
                    "</m:sSup>"
                )
            else:
                out.append(_math_text(token))
            continue  # мы уже обработали subscripts/superscripts сами

        # Сюда падаем только если сработала \text*-ветка — тогда продолжаем цикл
        continue

    return "".join(out)


def _read_token(expr: str, start: int) -> tuple[str, int]:
    ch = expr[start]
    if ch == "\\":
        j = start + 1
        while j < len(expr) and expr[j].isalpha():
            j += 1
        command = expr[start + 1:j]
        if command:
            return _LATEX_SYMBOLS.get(command, "\\" + command), j
        if j < len(expr):
            return expr[j], j + 1
        return "\\", j
    return ch, start + 1


def _read_script(expr: str, start: int) -> tuple[str, int]:
    if start < len(expr) and expr[start] == "{":
        return _read_group(expr, start)
    if start < len(expr):
        return expr[start], start + 1
    return "", start


def _read_group(expr: str, start: int) -> tuple[str, int]:
    while start < len(expr) and expr[start].isspace():
        start += 1
    if start >= len(expr) or expr[start] != "{":
        return "", start
    depth = 0
    for i in range(start, len(expr)):
        if expr[i] == "{":
            depth += 1
        elif expr[i] == "}":
            depth -= 1
            if depth == 0:
                return expr[start + 1:i], i + 1
    return expr[start + 1:], len(expr)


def _math_text(text: str) -> str:
    return f'<m:r><m:t xml:space="preserve">{_xml_escape(text)}</m:t></m:r>'


def _normalize_latex(latex: str) -> str:
    latex = latex.replace(r"\left", "").replace(r"\right", "")
    latex = latex.replace(r"\,", " ").replace(r"\;", " ").replace(r"\:", " ")
    latex = re.sub(r"\s+", " ", latex)
    return latex.strip()


def _xml_escape(text: Any) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


_LATEX_SYMBOLS = {
    "eta": "η",
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "theta": "θ",
    "sigma": "σ",
    "mu": "μ",
    "lambda": "λ",
    "Delta": "Δ",
    "times": "×",
    "cdot": "·",
    "leq": "≤",
    "geq": "≥",
    "neq": "≠",
    "approx": "≈",
    "in": "∈",
    "pm": "±",
    "%": "%",
}


def cards_to_tsv(cards: list[dict[str, Any]]) -> bytes:
    rows = []
    for card in cards:
        front = _clean_cell(card.get("front") or card.get("question") or "")
        back = _clean_cell(card.get("back") or card.get("answer") or "")
        source = _clean_cell(card.get("source_display") or card.get("source") or "")
        if front and back:
            rows.append(f"{front}\t{back} {source}".strip())
    return ("\n".join(rows) + "\n").encode("utf-8")


def cards_to_csv(cards: list[dict[str, Any]]) -> bytes:
    out = io.StringIO()
    writer = csv.writer(out, delimiter=";")
    writer.writerow(["Вопрос", "Ответ", "Источник"])
    for card in cards:
        front = _clean_cell(card.get("front") or card.get("question") or "")
        back = _clean_cell(card.get("back") or card.get("answer") or "")
        source = _clean_cell(card.get("source_display") or card.get("source") or "")
        if front and back:
            writer.writerow([front, back, source])
    return out.getvalue().encode("utf-8-sig")


def cards_docx_export(title: str, cards: list[dict[str, Any]]) -> bytes:
    document = docx.Document()
    document.add_heading(title.strip() or "Учебные карточки", level=1)
    for index, card in enumerate(cards, 1):
        front = _clean_cell(card.get("front") or card.get("question") or "")
        back = _clean_cell(card.get("back") or card.get("answer") or "")
        source = _clean_cell(card.get("source_display") or card.get("source") or "")
        if not front or not back:
            continue
        document.add_heading(f"Карточка {index}", level=2)
        q = document.add_paragraph()
        q.add_run("Вопрос: ").bold = True
        q.add_run(front)
        a = document.add_paragraph()
        a.add_run("Ответ: ").bold = True
        a.add_run(back)
        if source:
            s = document.add_paragraph()
            s.add_run("Источник: ").bold = True
            s.add_run(source)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def build_apkg(cards: list[dict[str, Any]], deck_name: str, *, package_id: str = "") -> bytes | None:
    try:
        import genanki
    except ImportError:
        return None

    deck_id = _stable_int(f"deck:{package_id}:{deck_name}")
    model_id = _stable_int("model:Навигатор Anki")
    model = genanki.Model(
        model_id,
        "Навигатор: вопрос-ответ",
        fields=[
            {"name": "Question"},
            {"name": "Answer"},
        ],
        templates=[{
            "name": "Card 1",
            "qfmt": "{{Question}}",
            "afmt": "{{FrontSide}}<hr id=\"answer\">{{Answer}}",
        }],
    )
    deck = genanki.Deck(deck_id, deck_name[:80] or "Навигатор")
    for card in cards:
        front = str(card.get("front") or card.get("question") or "").strip()
        back = str(card.get("back") or card.get("answer") or "").strip()
        source = str(card.get("source_display") or card.get("source") or "").strip()
        if not front or not back:
            continue
        note = genanki.Note(
            model=model,
            fields=[front, f"{back}<br><br><small>{source}</small>" if source else back],
            guid=hashlib.sha1(f"{package_id}|{front}|{back}".encode("utf-8")).hexdigest(),
        )
        deck.add_note(note)
    package = genanki.Package(deck)
    with tempfile.NamedTemporaryFile(suffix=".apkg", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        package.write_to_file(tmp_path)
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        try:
            import os
            os.remove(tmp_path)
        except OSError:
            pass


def save_flashcard_exports(
    cards: list[dict[str, Any]],
    deck_name: str,
    export_dir: str | Path,
    *,
    prefix: str = "navigator_flashcards",
    package_id: str = "",
) -> dict[str, Any]:
    export_path = Path(export_dir)
    export_path.mkdir(parents=True, exist_ok=True)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_prefix = safe_filename(prefix or deck_name or "navigator_flashcards")
    base_name = f"{stamp}_{safe_prefix}"

    tsv_bytes = cards_to_tsv(cards)
    tsv_path = export_path / f"{base_name}.tsv"
    tsv_path.write_bytes(tsv_bytes)

    csv_bytes = cards_to_csv(cards)
    csv_path = export_path / f"{base_name}.csv"
    csv_path.write_bytes(csv_bytes)

    docx_bytes = cards_docx_export(deck_name, cards)
    docx_path = export_path / f"{base_name}.docx"
    docx_path.write_bytes(docx_bytes)

    apkg_bytes = build_apkg(cards, deck_name, package_id=package_id)
    apkg_path = None
    if apkg_bytes:
        apkg_path = export_path / f"{base_name}.apkg"
        apkg_path.write_bytes(apkg_bytes)

    return {
        "tsv_bytes": tsv_bytes,
        "tsv_path": str(tsv_path),
        "csv_bytes": csv_bytes,
        "csv_path": str(csv_path),
        "docx_bytes": docx_bytes,
        "docx_path": str(docx_path),
        "apkg_bytes": apkg_bytes,
        "apkg_path": str(apkg_path) if apkg_path else "",
    }


def safe_filename(value: Any, *, max_len: int = 96) -> str:
    text = re.sub(r"\s+", "_", str(value or "").strip())
    text = re.sub(r"[^\w.\-]+", "_", text, flags=re.UNICODE).strip("._-")
    return (text[:max_len].strip("._-") or "navigator_flashcards")


def graphviz_dot(graph: dict[str, Any]) -> str:
    nodes = graph.get("nodes") or []
    edges = graph.get("edges") or []
    lines = ["graph G {", "  graph [rankdir=LR, bgcolor=\"transparent\"];", "  node [shape=box, style=\"rounded,filled\", fillcolor=\"#111111\", color=\"#2a2a2a\", fontcolor=\"#fafafa\"];", "  edge [color=\"#525252\", fontcolor=\"#a3a3a3\"];"]
    for node in nodes:
        node_id = _dot_id(node.get("id") or node.get("label") or "")
        label = str(node.get("label") or node.get("id") or "").replace('"', '\\"')
        if node_id:
            lines.append(f'  "{node_id}" [label="{label}"];')
    for edge in edges:
        source = _dot_id(edge.get("source") or "")
        target = _dot_id(edge.get("target") or "")
        label = str(edge.get("label") or "").replace('"', '\\"')
        if source and target:
            suffix = f' [label="{label}"]' if label else ""
            lines.append(f'  "{source}" -- "{target}"{suffix};')
    lines.append("}")
    return "\n".join(lines)


def _stable_int(seed: str) -> int:
    return int(hashlib.sha1(seed.encode("utf-8")).hexdigest()[:8], 16)


def _clean_cell(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().replace("\t", " ")


def _dot_id(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()
