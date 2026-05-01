from __future__ import annotations

import hashlib
import io
import json
import re
import tempfile
import csv
from datetime import datetime
from pathlib import Path
from typing import Any

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml


def fragments_context(fragments: list[dict[str, Any]], *, max_chars: int = 12000) -> str:
    parts = []
    used = 0
    for index, fragment in enumerate(fragments, 1):
        text = str(fragment.get("text", "")).strip()
        if not text:
            continue
        header = f"[{index}] Документ: {fragment.get('document', '')}, стр. {fragment.get('page', '')}\n"
        block = header + text + "\n\n"
        if used + len(block) > max_chars and parts:
            break
        parts.append(block)
        used += len(block)
    return "".join(parts)


def parse_json_loose(raw: str) -> dict[str, Any]:
    text = (raw or "").strip()
    text = re.sub(r"^```(?:json)?\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"\s*```$", "", text)
    try:
        data = json.loads(text)
        return data if isinstance(data, dict) else {"items": data}
    except json.JSONDecodeError:
        pass

    start = text.find("{")
    end = text.rfind("}")
    if start >= 0 and end > start:
        data = json.loads(text[start:end + 1])
        return data if isinstance(data, dict) else {"items": data}
    raise ValueError("LLM вернула невалидный JSON.")


def markdown_export(title: str, body: str, fragments: list[dict[str, Any]] | None = None) -> bytes:
    lines = [f"# {title.strip() or 'Экспорт'}", "", body.strip(), ""]
    if fragments:
        lines.extend(["## Источники", ""])
        for index, fragment in enumerate(fragments, 1):
            lines.append(
                f"- [{index}] {fragment.get('document', '')}, стр. {fragment.get('page', '')}"
            )
    return ("\n".join(lines).strip() + "\n").encode("utf-8")


def docx_export(title: str, body: str, fragments: list[dict[str, Any]] | None = None) -> bytes:
    document = docx.Document()
    document.add_heading(title.strip() or "Экспорт", level=1)
    for part in re.split(r"(\$\$.*?\$\$)", body.strip(), flags=re.DOTALL):
        part = part.strip()
        if not part:
            continue
        if part.startswith("$$") and part.endswith("$$"):
            latex = part[2:-2].strip()
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph._p.append(_omath(latex))
            continue
        for block in re.split(r"\n{2,}", part):
            block = block.strip()
            if not block:
                continue
            if block.startswith(("# ", "## ", "### ")):
                level = min(block.count("#", 0, block.find(" ")), 3)
                document.add_heading(block.lstrip("# ").strip(), level=level)
            else:
                paragraph = document.add_paragraph()
                _append_inline_math(paragraph, block)

    if fragments:
        document.add_heading("Источники", level=2)
        for index, fragment in enumerate(fragments, 1):
            document.add_paragraph(
                f"[{index}] {fragment.get('document', '')}, стр. {fragment.get('page', '')}"
            )

    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def _append_inline_math(paragraph: Any, text: str) -> None:
    for part in re.split(r"(\$[^$\n]+\$)", text):
        if not part:
            continue
        if part.startswith("$") and part.endswith("$"):
            paragraph._p.append(_omath(part[1:-1].strip()))
        else:
            paragraph.add_run(part)


def _omath(latex: str) -> Any:
    xml = (
        '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        + _math_fragments(_normalize_latex(latex)) +
        "</m:oMath>"
    )
    return parse_xml(xml)


def _math_fragments(expr: str) -> str:
    out = []
    i = 0
    while i < len(expr):
        if expr.startswith(r"\frac", i):
            i += len(r"\frac")
            num, i = _read_group(expr, i)
            den, i = _read_group(expr, i)
            out.append(
                "<m:f><m:fPr><m:type m:val=\"bar\"/></m:fPr>"
                f"<m:num>{_math_fragments(num)}</m:num>"
                f"<m:den>{_math_fragments(den)}</m:den>"
                "</m:f>"
            )
            continue

        token, i = _read_token(expr, i)
        if not token:
            break
        sub = sup = None
        while i < len(expr) and expr[i] in "_^":
            marker = expr[i]
            value, i = _read_script(expr, i + 1)
            if marker == "_":
                sub = value
            else:
                sup = value
        if sub is not None and sup is not None:
            out.append(
                "<m:sSubSup>"
                f"<m:e>{_math_text(token)}</m:e>"
                f"<m:sub>{_math_fragments(sub)}</m:sub>"
                f"<m:sup>{_math_fragments(sup)}</m:sup>"
                "</m:sSubSup>"
            )
        elif sub is not None:
            out.append(
                "<m:sSub>"
                f"<m:e>{_math_text(token)}</m:e>"
                f"<m:sub>{_math_fragments(sub)}</m:sub>"
                "</m:sSub>"
            )
        elif sup is not None:
            out.append(
                "<m:sSup>"
                f"<m:e>{_math_text(token)}</m:e>"
                f"<m:sup>{_math_fragments(sup)}</m:sup>"
                "</m:sSup>"
            )
        else:
            out.append(_math_text(token))
    return "".join(out)


def _read_token(expr: str, start: int) -> tuple[str, int]:
    ch = expr[start]
    if ch == "\\":
        j = start + 1
        while j < len(expr) and expr[j].isalpha():
            j += 1
        command = expr[start + 1:j]
        if command:
            return _LATEX_SYMBOLS.get(command, "\\" + command), j
        if j < len(expr):
            return expr[j], j + 1
        return "\\", j
    return ch, start + 1


def _read_script(expr: str, start: int) -> tuple[str, int]:
    if start < len(expr) and expr[start] == "{":
        return _read_group(expr, start)
    if start < len(expr):
        return expr[start], start + 1
    return "", start


def _read_group(expr: str, start: int) -> tuple[str, int]:
    while start < len(expr) and expr[start].isspace():
        start += 1
    if start >= len(expr) or expr[start] != "{":
        return "", start
    depth = 0
    for i in range(start, len(expr)):
        if expr[i] == "{":
            depth += 1
        elif expr[i] == "}":
            depth -= 1
            if depth == 0:
                return expr[start + 1:i], i + 1
    return expr[start + 1:], len(expr)


def _math_text(text: str) -> str:
    return f'<m:r><m:t xml:space="preserve">{_xml_escape(text)}</m:t></m:r>'


def _normalize_latex(latex: str) -> str:
    latex = latex.replace(r"\left", "").replace(r"\right", "")
    latex = latex.replace(r"\,", " ").replace(r"\;", " ").replace(r"\:", " ")
    latex = re.sub(r"\s+", " ", latex)
    return latex.strip()


def _xml_escape(text: Any) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


_LATEX_SYMBOLS = {
    "eta": "η",
    "alpha": "α",
    "beta": "β",
    "gamma": "γ",
    "theta": "θ",
    "sigma": "σ",
    "mu": "μ",
    "lambda": "λ",
    "Delta": "Δ",
    "times": "×",
    "cdot": "·",
    "leq": "≤",
    "geq": "≥",
    "neq": "≠",
    "approx": "≈",
    "in": "∈",
    "pm": "±",
    "%": "%",
}


def cards_to_tsv(cards: list[dict[str, Any]]) -> bytes:
    rows = []
    for card in cards:
        front = _clean_cell(card.get("front") or card.get("question") or "")
        back = _clean_cell(card.get("back") or card.get("answer") or "")
        source = _clean_cell(card.get("source_display") or card.get("source") or "")
        if front and back:
            rows.append(f"{front}\t{back} {source}".strip())
    return ("\n".join(rows) + "\n").encode("utf-8")


def cards_to_csv(cards: list[dict[str, Any]]) -> bytes:
    out = io.StringIO()
    writer = csv.writer(out, delimiter=";")
    writer.writerow(["Вопрос", "Ответ", "Источник"])
    for card in cards:
        front = _clean_cell(card.get("front") or card.get("question") or "")
        back = _clean_cell(card.get("back") or card.get("answer") or "")
        source = _clean_cell(card.get("source_display") or card.get("source") or "")
        if front and back:
            writer.writerow([front, back, source])
    return out.getvalue().encode("utf-8-sig")


def cards_docx_export(title: str, cards: list[dict[str, Any]]) -> bytes:
    document = docx.Document()
    document.add_heading(title.strip() or "Учебные карточки", level=1)
    for index, card in enumerate(cards, 1):
        front = _clean_cell(card.get("front") or card.get("question") or "")
        back = _clean_cell(card.get("back") or card.get("answer") or "")
        source = _clean_cell(card.get("source_display") or card.get("source") or "")
        if not front or not back:
            continue
        document.add_heading(f"Карточка {index}", level=2)
        q = document.add_paragraph()
        q.add_run("Вопрос: ").bold = True
        q.add_run(front)
        a = document.add_paragraph()
        a.add_run("Ответ: ").bold = True
        a.add_run(back)
        if source:
            s = document.add_paragraph()
            s.add_run("Источник: ").bold = True
            s.add_run(source)
    out = io.BytesIO()
    document.save(out)
    return out.getvalue()


def build_apkg(cards: list[dict[str, Any]], deck_name: str, *, package_id: str = "") -> bytes | None:
    try:
        import genanki
    except ImportError:
        return None

    deck_id = _stable_int(f"deck:{package_id}:{deck_name}")
    model_id = _stable_int("model:Навигатор Anki")
    model = genanki.Model(
        model_id,
        "Навигатор: вопрос-ответ",
        fields=[
            {"name": "Question"},
            {"name": "Answer"},
        ],
        templates=[{
            "name": "Card 1",
            "qfmt": "{{Question}}",
            "afmt": "{{FrontSide}}<hr id=\"answer\">{{Answer}}",
        }],
    )
    deck = genanki.Deck(deck_id, deck_name[:80] or "Навигатор")
    for card in cards:
        front = str(card.get("front") or card.get("question") or "").strip()
        back = str(card.get("back") or card.get("answer") or "").strip()
        source = str(card.get("source_display") or card.get("source") or "").strip()
        if not front or not back:
            continue
        note = genanki.Note(
            model=model,
            fields=[front, f"{back}<br><br><small>{source}</small>" if source else back],
            guid=hashlib.sha1(f"{package_id}|{front}|{back}".encode("utf-8")).hexdigest(),
        )
        deck.add_note(note)
    package = genanki.Package(deck)
    with tempfile.NamedTemporaryFile(suffix=".apkg", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        package.write_to_file(tmp_path)
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        try:
            import os
            os.remove(tmp_path)
        except OSError:
            pass


def save_flashcard_exports(
    cards: list[dict[str, Any]],
    deck_name: str,
    export_dir: str | Path,
    *,
    prefix: str = "navigator_flashcards",
    package_id: str = "",
) -> dict[str, Any]:
    export_path = Path(export_dir)
    export_path.mkdir(parents=True, exist_ok=True)

    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_prefix = safe_filename(prefix or deck_name or "navigator_flashcards")
    base_name = f"{stamp}_{safe_prefix}"

    tsv_bytes = cards_to_tsv(cards)
    tsv_path = export_path / f"{base_name}.tsv"
    tsv_path.write_bytes(tsv_bytes)

    csv_bytes = cards_to_csv(cards)
    csv_path = export_path / f"{base_name}.csv"
    csv_path.write_bytes(csv_bytes)

    docx_bytes = cards_docx_export(deck_name, cards)
    docx_path = export_path / f"{base_name}.docx"
    docx_path.write_bytes(docx_bytes)

    apkg_bytes = build_apkg(cards, deck_name, package_id=package_id)
    apkg_path = None
    if apkg_bytes:
        apkg_path = export_path / f"{base_name}.apkg"
        apkg_path.write_bytes(apkg_bytes)

    return {
        "tsv_bytes": tsv_bytes,
        "tsv_path": str(tsv_path),
        "csv_bytes": csv_bytes,
        "csv_path": str(csv_path),
        "docx_bytes": docx_bytes,
        "docx_path": str(docx_path),
        "apkg_bytes": apkg_bytes,
        "apkg_path": str(apkg_path) if apkg_path else "",
    }


def safe_filename(value: Any, *, max_len: int = 96) -> str:
    text = re.sub(r"\s+", "_", str(value or "").strip())
    text = re.sub(r"[^\w.\-]+", "_", text, flags=re.UNICODE).strip("._-")
    return (text[:max_len].strip("._-") or "navigator_flashcards")


def graphviz_dot(graph: dict[str, Any]) -> str:
    nodes = graph.get("nodes") or []
    edges = graph.get("edges") or []
    lines = ["graph G {", "  graph [rankdir=LR, bgcolor=\"transparent\"];", "  node [shape=box, style=\"rounded,filled\", fillcolor=\"#111111\", color=\"#2a2a2a\", fontcolor=\"#fafafa\"];", "  edge [color=\"#525252\", fontcolor=\"#a3a3a3\"];"]
    for node in nodes:
        node_id = _dot_id(node.get("id") or node.get("label") or "")
        label = str(node.get("label") or node.get("id") or "").replace('"', '\\"')
        if node_id:
            lines.append(f'  "{node_id}" [label="{label}"];')
    for edge in edges:
        source = _dot_id(edge.get("source") or "")
        target = _dot_id(edge.get("target") or "")
        label = str(edge.get("label") or "").replace('"', '\\"')
        if source and target:
            suffix = f' [label="{label}"]' if label else ""
            lines.append(f'  "{source}" -- "{target}"{suffix};')
    lines.append("}")
    return "\n".join(lines)


def _stable_int(seed: str) -> int:
    return int(hashlib.sha1(seed.encode("utf-8")).hexdigest()[:8], 16)


def _clean_cell(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip().replace("\t", " ")


def _dot_id(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()
